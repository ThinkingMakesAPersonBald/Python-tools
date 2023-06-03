import lxml
import docx
import os


def get_word_count(docx_file_path):
    doc = docx.Document(docx_file_path)
    word_count = 0

    for paragraph in doc.paragraphs:
        words = paragraph.text.split()
        word_count += len(words)

    return word_count
distinct_words = set()
def get_distinct_word_count(docx_file_path):
    doc = docx.Document(docx_file_path)
    

    for paragraph in doc.paragraphs:
        words = paragraph.text.split()
        for word in words:
            distinct_words.add(word)

    return len(distinct_words)
# file_path = '/Users/peixinhua/Downloads/Count the words'
# word_count = get_distinct_word_count(os.path.join(file_path,'New Concept English Four.docx'))
# print(f'The word count is: {word_count}')

root_path = '/Users/peixinhua/Downloads/Count the words'
# Walk through all the directores and files in the directory tree

for root, dirs, files in os.walk(root_path):
    for file_name in files:
        word_count = get_distinct_word_count(os.path.join(root_path,file_name))
        print(file_name,word_count)

